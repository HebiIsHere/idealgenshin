"""将使用说明 markdown 转换为格式化的 docx"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\zxy\Desktop\理想原生v4.3.0-使用说明.md"
DST = r"C:\Users\zxy\Desktop\理想原生v4.3.0-使用说明.docx"

# 芙宁娜水蓝色
HYDRO_MID = RGBColor(0x5B, 0xC0, 0xEB)
HYDRO_DEEP = RGBColor(0x3A, 0xA0, 0xC8)
CELADON_MID = RGBColor(0xD0, 0xE4, 0xDC)
DARK_BG = RGBColor(0x0B, 0x14, 0x24)
TEXT_PRIMARY = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_SECONDARY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_COLOR = "5BC0EB"
LIGHT_BLUE_BG = "E8F6FC"
ACCENT_GOLD = RGBColor(0xC8, 0xA8, 0x50)

FONT_TITLE = "Noto Serif SC"
FONT_BODY = "Noto Sans SC"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = Pt(10.5)
style.font.color.rgb = TEXT_PRIMARY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
# Set East Asian font
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), FONT_BODY)
rPr.insert(0, rFonts)

for level, (size, color, spacing_before, spacing_after) in {
    1: (22, HYDRO_DEEP, Pt(36), Pt(8)),
    2: (16, HYDRO_DEEP, Pt(28), Pt(6)),
    3: (13, HYDRO_MID, Pt(20), Pt(4)),
}.items():
    s = doc.styles[f'Heading {level}']
    s.font.name = FONT_TITLE
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = spacing_before
    s.paragraph_format.space_after = spacing_after
    rPr = s.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_TITLE)
    rPr.insert(0, rFonts)

# ── Helper functions ──

def add_paragraph_styled(text, bold=False, italic=False, color=None, size=None, alignment=None, font_name=None):
    """Add a paragraph with inline formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.insert(0, rFonts)
    return p

def add_paragraph_with_inline(segments):
    """segments is a list of (text, bold, italic, is_code) tuples."""
    p = doc.add_paragraph()
    for text, bold, italic, is_code in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if is_code:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        if not text.strip():
            continue
    return p

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        if header:
            run.bold = True
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, "3AA0C8")
        else:
            run.font.color.rgb = TEXT_PRIMARY
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_BODY)
        rPr.insert(0, rFonts)
    return row

def add_horizontal_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Parse markdown ──

with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.readlines()

def parse_inline(text):
    """Parse bold, italic, inline code from a line of text. Returns list of (text, bold, italic, is_code)."""
    segments = []
    # Handle bold+italic combined first, then bold, then italic, then code
    # Bold: **text**
    # Code: `text`
    # Simple regex-based parser
    pattern = r'(\*\*(.+?)\*\*|`([^`]+)`|(.+?))'
    remaining = text
    while remaining:
        # Find bold **...**
        m = re.match(r'\*\*(.+?)\*\*', remaining)
        if m:
            segments.append((m.group(1), True, False, False))
            remaining = remaining[m.end():]
            continue
        # Find inline code `...`
        m = re.match(r'`([^`]+)`', remaining)
        if m:
            segments.append((m.group(1), False, False, True))
            remaining = remaining[m.end():]
            continue
        # Find italic *...* (single asterisk, not double)
        m = re.match(r'(?<!\*)\*([^*]+)\*(?!\*)', remaining)
        if m:
            segments.append((m.group(1), False, True, False))
            remaining = remaining[m.end():]
            continue
        # Plain text up to next special char
        m = re.match(r'([^*`]+)', remaining)
        if m:
            segments.append((m.group(1), False, False, False))
            remaining = remaining[m.end():]
            continue
        # Single character
        segments.append((remaining[0], False, False, False))
        remaining = remaining[1:]
    return segments

i = 0
while i < len(lines):
    line = lines[i]

    # Skip empty lines
    if not line.strip():
        i += 1
        continue

    # Horizontal rule
    if line.strip() == '---':
        add_horizontal_rule()
        i += 1
        continue

    # Headers
    h3_match = re.match(r'^### (.+)$', line)
    h2_match = re.match(r'^## (.+)$', line)
    h1_match = re.match(r'^# (.+)$', line)

    if h3_match:
        doc.add_heading(h3_match.group(1), level=3)
        i += 1
        continue
    if h2_match:
        doc.add_heading(h2_match.group(1), level=2)
        i += 1
        continue
    if h1_match:
        doc.add_heading(h1_match.group(1), level=1)
        i += 1
        continue

    # Blockquote
    if line.startswith('> '):
        bq_lines = []
        while i < len(lines) and lines[i].startswith('> '):
            bq_lines.append(lines[i][2:].rstrip())
            i += 1
        bq_text = ' '.join(bq_lines)
        bq_para = doc.add_paragraph()
        bq_para.paragraph_format.left_indent = Cm(1.0)
        bq_run = bq_para.add_run(bq_text)
        bq_run.font.name = FONT_BODY
        bq_run.font.size = Pt(11)
        bq_run.font.color.rgb = HYDRO_DEEP
        bq_run.italic = True
        rPr = bq_run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_BODY)
        rPr.insert(0, rFonts)
        # Add left border shading via paragraph indent +  shading
        pPr = bq_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), '5BC0EB')
        pBdr.append(left)
        pPr.append(pBdr)
        continue

    # Unordered list
    ul_match = re.match(r'^- (.+)$', line)
    if ul_match:
        ul_lines = []
        while i < len(lines) and re.match(r'^- (.+)$', lines[i]):
            ul_lines.append(re.match(r'^- (.+)$', lines[i]).group(1))
            i += 1
        for item in ul_lines:
            p = doc.add_paragraph(style='List Bullet')
            # Clear default run and add formatted
            p.clear()
            segs = parse_inline(item)
            for text, bold, italic, is_code in segs:
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                if is_code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
        continue

    # Ordered list
    ol_match = re.match(r'^\d+\. (.+)$', line)
    if ol_match:
        ol_lines = []
        while i < len(lines) and re.match(r'^\d+\. (.+)$', lines[i]):
            ol_lines.append(re.match(r'^\d+\. (.+)$', lines[i]).group(1))
            i += 1
        for item in ol_lines:
            p = doc.add_paragraph(style='List Number')
            p.clear()
            segs = parse_inline(item)
            for text, bold, italic, is_code in segs:
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                if is_code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
        continue

    # Table detection: look ahead for separator line containing | and ---
    if '|' in line:
        # Check if next line is a separator
        if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
            header_line = line
            sep_line = lines[i + 1]
            i += 2
            # Collect data rows
            data_rows = []
            while i < len(lines) and '|' in lines[i]:
                data_rows.append(lines[i])
                i += 1

            # Parse header
            headers = [c.strip() for c in header_line.strip().strip('|').split('|')]
            # Parse alignment from separator
            seps = [c.strip() for c in sep_line.strip().strip('|').split('|')]
            alignments = []
            for s in seps:
                if s.startswith(':') and s.endswith(':'):
                    alignments.append('center')
                elif s.endswith(':'):
                    alignments.append('right')
                else:
                    alignments.append('left')

            # Parse data
            parsed_rows = []
            for row_line in data_rows:
                parsed_rows.append([c.strip() for c in row_line.strip().strip('|').split('|')])

            # Create table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            hdr = table.rows[0]
            for j, h in enumerate(headers):
                cell = hdr.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.name = FONT_BODY
                run.font.size = Pt(10)
                run.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), FONT_BODY)
                rPr.insert(0, rFonts)
                set_cell_shading(cell, "3AA0C8")

            # Data rows
            for row_data in parsed_rows:
                row = table.add_row()
                for j, val in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(val)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)
                    run.font.color.rgb = TEXT_PRIMARY
                    rPr = run._element.get_or_add_rPr()
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:eastAsia'), FONT_BODY)
                    rPr.insert(0, rFonts)
                    if j < len(alignments):
                        if alignments[j] == 'center':
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif alignments[j] == 'right':
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add spacing after table
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(4)
            p_space.paragraph_format.space_after = Pt(4)
            continue

    # Regular paragraph with inline formatting
    segs = parse_inline(line.rstrip())
    if segs:
        p = doc.add_paragraph()
        for text, bold, italic, is_code in segs:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            if is_code:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)

    i += 1

# ── Save ──
doc.save(DST)
print(f"Done → {DST}")
