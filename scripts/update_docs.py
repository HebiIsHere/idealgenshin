"""更新理想原生项目中的 docx 文档版本号和内容。

用法:
    python update_docs.py <旧版本号> <新版本号>

示例:
    python update_docs.py v4.2.1 v4.2.2
"""
import sys, os, re
from docx import Document

PUBLIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'public')

DOC_FILES = [
    '理想原生{old}-使用说明.docx',
    '理想原生{old}-项目介绍手册.docx',
]


def replace_in_runs(paragraphs, old: str, new: str, old_dot: str, new_dot: str) -> int:
    """Replace version strings in all runs across paragraphs. Returns change count."""
    count = 0
    for p in paragraphs:
        for run in p.runs:
            text = run.text
            if old in text or old_dot in text:
                run.text = text.replace(old, new)
                # dot-version: only replace standalone version numbers, not dates
                run.text = re.sub(rf'\b{re.escape(old_dot)}\b', new_dot, run.text)
                count += 1
    return count


def main():
    if len(sys.argv) != 3:
        print(f'用法: python {sys.argv[0]} <旧版本号> <新版本号>')
        sys.exit(1)

    old = sys.argv[1]
    new = sys.argv[2]
    # Build dot-versions for matching (v4.1 → 4.1, v4.2.1 → 4.2.1)
    old_dot = old.lstrip('v')
    new_dot = new.lstrip('v')

    for fname_template in DOC_FILES:
        old_fname = fname_template.format(old=old)
        new_fname = fname_template.format(old=new)
        src = os.path.join(PUBLIC_DIR, old_fname)
        dst = os.path.join(PUBLIC_DIR, new_fname)

        if not os.path.exists(src):
            print(f'[SKIP] {old_fname} 不存在')
            continue

        doc = Document(src)
        changes = 0

        # Update main body paragraphs
        changes += replace_in_runs(doc.paragraphs, old, new, old_dot, new_dot)

        # Update tables
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    changes += replace_in_runs(c.paragraphs, old, new, old_dot, new_dot)

        doc.save(dst)
        os.remove(src)
        print(f'[OK] {old_fname} → {new_fname} ({changes} 处替换)')

    print('完成。别忘更新 LandingPage.tsx 中的 href 链接。')


if __name__ == '__main__':
    main()
